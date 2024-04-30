import os
import re

import docx
import pysrt

from const import dash_pattern
from tt.transformer import transformer_2

MAX_LENGTH = 150
MIN_PARAGRAPH_WORDS = 100

sent_pattern = r"[!.?]"
combine_whitespace = re.compile(r"(?a:\s+)")

msg_by_lang = {
    'he': 'תמליל אוטומאטי שנוצר על ידי מכונה. יש לקחת בחשבון שיכולים להיות אי דיוקים בטקסט.',
    'en': '',
    'ru': '',
}


def srt_to_paragraphs(path, lang):
    subs = pysrt.open(path)
    res = ""
    for sub in subs:
        res += f"{sub.text} "
    res = combine_whitespace.sub(" ", res)
    words = res.split(" ")
    i = 0
    prev_sent_pos = 0
    paragraphs = []
    for w in words:
        if lang == "he" and re.search(dash_pattern, w) is not None and prev_sent_pos != i:
            p = " ".join(words[prev_sent_pos: i])
            paragraphs.append(p)
            prev_sent_pos = i
        if re.search(sent_pattern, w) is not None and i - prev_sent_pos > MIN_PARAGRAPH_WORDS:
            p = " ".join(words[prev_sent_pos: i + 1])
            paragraphs.append(p)
            prev_sent_pos = i + 1
        i += 1
    return paragraphs


def save_to_docx(paragraphs, lang, path):
    if lang == 'he':
        tpl = os.path.join(os.path.dirname(__file__), 'tpl_rtl.docx')
    else:
        tpl = os.path.join(os.path.dirname(__file__), 'tpl_ltr.docx')
    doc = docx.Document(tpl)
    add_paragraph(doc, msg_by_lang[lang], style='h3')
    for txt in paragraphs:
        add_paragraph(doc, txt)
    doc.save(path)


def add_paragraph(doc, txt, style='rtl'):
    p = doc.add_paragraph(style=style)
    r = p.add_run()
    font = r.font
    font.complex_script = True
    font.rtl = True
    r.add_text(txt)


def fix_text(txt):
    return transformer_2(txt)
